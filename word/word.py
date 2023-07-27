from docxtpl import DocxTemplate
import docx
import os
from copy import copy


class FillTemplate:
    def __init__(self, group, template_path='template'):
        self.group = group
        self.template_path = template_path
        self.doc = DocxTemplate('template.docx')
        self.out_path = 'output'
        self.global_context = {
            'diplomaDate': self.group.diplomaDate,
            'orientation': self.group.orientation,
            'qualification': self.group.qualification,
            'direction': self.group.direction,
            'vice_rector': self.group.vice_rector,
            'code': self.group.code,
            'faculty': self.group.faculty,
            'protocol': '({})'.format(self.group.protocol),
        }

    def fill_words(self):
        qualification_dict = {
            'Бакалавриат': 'бакалавра',
            'Специалитет': 'специалиста',
            'Магистратура': 'магистра',
            'Аспирантура': 'аспиранта'}
        try:
            os.mkdir(self.out_path)
        except Exception as e:
            print(e)
        try:
            os.mkdir('{}/{}'.format(self.out_path, self.group.name))
        except FileExistsError:
            pass

        for student in self.group.students:
            j = 1

            context = copy(self.global_context)

            context['qualification_with_mark'] = qualification_dict[self.group.qualification] + (' с отличием' * student.excel)
            context['first_name'] = student.firstName
            context['second_name'] = student.secondName
            context['third_name'] = student.thirdName
            context['birth_date'] = student.birthDate
            context['document'] = student.document
            context['certificateYear'] = student.certificateYear
            context['student_number'] = student.studentNumber

            if student.formCheck:
                context['info{}'.format(j)] = 'Форма обучения: {}'.format(self.group.form.lower())
                j += 1

            if student.seqFormCheck:
                context['info{}'.format(j)] = 'Сочетание форм обучения:'
                j += 1

            if student.speedrunCheck:
                context['info{}'.format(j)] = 'Пройдено ускоренное обучение по образовательной программе'
                j += 1

            if student.imposterCheck:
                context['info{}'.format(j)] = 'Часть образовательной программы в объеме з.е. освоена в'

            for i, exam in enumerate(student.first_page_exams):
                context['exam{}'.format(i)] = exam[0]
                context['exam{}_unit'.format(i)] = exam[1]
                context['exam{}_mark'.format(i)] = exam[2]

            for i, exam in enumerate(student.second_page_exams):
                context['exam_new{}'.format(i)] = exam[0]
                context['exam_new{}_unit'.format(i)] = exam[1]
                context['exam_new{}_mark'.format(i)] = exam[2]

            self.doc.render(context=context)
            self.save('{} {} {}'.format(student.secondName, student.firstName, student.thirdName))

    def save(self, filename):
        save_path = os.path.join(self.out_path, self.group.name, '{}.docx'.format(filename))
        self.doc.save(save_path)
        doc = docx.Document(save_path)
        first_table = doc.tables[-1]._cells[0].tables[0]
        second_table = doc.tables[-1]._cells[2].tables[0]

        def remove_row(table, row):
            tbl = table._tbl
            tr = row._tr
            tbl.remove(tr)

        def delete_rows(table):
            while not table.rows[-1].cells[0].text:
                remove_row(table, table.rows[-1])

        delete_rows(first_table)
        delete_rows(second_table)

        doc.save(save_path)


